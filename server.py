from __future__ import annotations

import json
import os
import shutil
import sqlite3
import csv
import threading
import time
import base64
import mimetypes
import traceback
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

import httpx
from dotenv import load_dotenv
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from mcp.server.fastmcp import FastMCP
from openai import OpenAI
from openpyxl import Workbook, load_workbook
from docx import Document
from pypdf import PdfReader
from pypdfium2 import PdfDocument

load_dotenv()

DB_PATH = Path(os.getenv("MCP_DEMO_DB_PATH", "demo.db"))
MCP_TRANSPORT = os.getenv("MCP_TRANSPORT", "stdio")
MCP_HOST = os.getenv("MCP_HOST", "127.0.0.1")
MCP_PORT = int(os.getenv("MCP_PORT", "8000"))
MCP_PATH = os.getenv("MCP_PATH", "/mcp")
HEARTBEAT_SECONDS = int(os.getenv("MCP_HEARTBEAT_SECONDS", "30"))
FILE_OPS_ROOT = os.getenv("MCP_FILE_OPS_ROOT")

mcp = FastMCP(
    "local-mcp-demo",
    host=MCP_HOST,
    port=MCP_PORT,
    streamable_http_path=MCP_PATH,
)


def _db_connection() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def _bootstrap_db() -> None:
    with _db_connection() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS notes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                body TEXT NOT NULL,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        count = conn.execute("SELECT COUNT(*) AS c FROM notes").fetchone()["c"]
        if count == 0:
            conn.executemany(
                "INSERT INTO notes(title, body) VALUES(?, ?)",
                [
                    ("Welcome", "Your local MCP SQLite tool is working."),
                    ("Next Step", "Try running SELECT * FROM notes;"),
                ],
            )


def _resolve_file_ops_path(path: str | None = None) -> Path:
    if not FILE_OPS_ROOT:
        raise ValueError("MCP_FILE_OPS_ROOT is not configured in .env.")

    root = Path(FILE_OPS_ROOT).expanduser().resolve()
    root.mkdir(parents=True, exist_ok=True)

    target = root if path is None else (root / path).resolve()
    if target != root and root not in target.parents:
        raise ValueError("Path escapes the configured MCP_FILE_OPS_ROOT.")
    return target


SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
SUPPORTED_SPREADSHEET_EXTENSIONS = {".csv", ".tsv", ".xlsx"}
GOOGLE_CALENDAR_SCOPES = ["https://www.googleapis.com/auth/calendar.events"]
CALENDAR_TIMEZONE = ZoneInfo("America/New_York")


LOG_PATH = Path(os.getenv("MCP_TOOL_LOG_PATH", "mcp_tool_events.jsonl"))


def _write_tool_event(event: dict[str, Any]) -> None:
    event.setdefault("timestamp", datetime.now(timezone.utc).isoformat())
    serialized = json.dumps(event, ensure_ascii=False)
    LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
    with LOG_PATH.open("a", encoding="utf-8") as handle:
        handle.write(serialized + "\n")
    print(f"[mcp-tool-log] {serialized}", flush=True)


def _truncate_for_log(value: Any, max_chars: int = 2000) -> Any:
    try:
        as_text = json.dumps(value, ensure_ascii=False)
    except TypeError:
        as_text = str(value)
    if len(as_text) <= max_chars:
        return value
    return f"{as_text[:max_chars]}...<truncated>"


def _run_tool_with_logging(tool_name: str, tool_args: dict[str, Any], fn: Any) -> str:
    start = time.time()
    _write_tool_event({"event": "tool_start", "tool": tool_name, "args": _truncate_for_log(tool_args)})
    try:
        result = fn()
        _write_tool_event(
            {
                "event": "tool_success",
                "tool": tool_name,
                "duration_ms": int((time.time() - start) * 1000),
                "result_preview": _truncate_for_log(result),
            }
        )
        return result
    except Exception as exc:
        _write_tool_event(
            {
                "event": "tool_error",
                "tool": tool_name,
                "duration_ms": int((time.time() - start) * 1000),
                "error_type": type(exc).__name__,
                "error": str(exc),
                "traceback": traceback.format_exc(),
            }
        )
        raise


def _get_google_calendar_service() -> Any:
    creds_env = os.getenv("GOOGLE_CALENDAR_CREDENTIALS_PATH")
    if creds_env:
        creds_path = Path(creds_env).expanduser()
    else:
        default_path = Path("google_credentials.json")
        if default_path.exists():
            creds_path = default_path
        else:
            matches = sorted(Path(".").glob("client_secret*.json"))
            creds_path = matches[0] if matches else default_path
    token_path = Path(os.getenv("GOOGLE_CALENDAR_TOKEN_PATH", "google_token.json")).expanduser()

    creds: Credentials | None = None
    if token_path.exists():
        creds = Credentials.from_authorized_user_file(str(token_path), GOOGLE_CALENDAR_SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not creds_path.exists():
                raise ValueError(
                    f"Google credentials file not found: {creds_path}. "
                    "Set GOOGLE_CALENDAR_CREDENTIALS_PATH in .env to your downloaded client_secret JSON file."
                )
            flow = InstalledAppFlow.from_client_secrets_file(str(creds_path), GOOGLE_CALENDAR_SCOPES)
            creds = flow.run_local_server(port=0)
        token_path.parent.mkdir(parents=True, exist_ok=True)
        token_path.write_text(creds.to_json(), encoding="utf-8")

    return build("calendar", "v3", credentials=creds)


def get_now_et() -> datetime:
    return datetime.now(CALENDAR_TIMEZONE)


def get_default_current_week_window_et(now_et: datetime) -> tuple[datetime, datetime]:
    end_of_week_date = (now_et + timedelta(days=(6 - now_et.weekday()))).date()
    end_of_week = datetime.combine(end_of_week_date, datetime.max.time(), tzinfo=CALENDAR_TIMEZONE)
    end_of_week = end_of_week.replace(microsecond=999000)
    return now_et, end_of_week


def _parse_window_bound_to_et(value: str) -> datetime:
    normalized = value.strip().lower()
    now_et = get_now_et()
    if normalized in {"upcoming", "this week"}:
        start, _ = get_default_current_week_window_et(now_et)
        return start
    if normalized == "tomorrow":
        tomorrow = now_et + timedelta(days=1)
        return tomorrow.replace(hour=0, minute=0, second=0, microsecond=0)
    if normalized == "next week":
        next_monday_date = (now_et + timedelta(days=(7 - now_et.weekday()))).date()
        return datetime.combine(next_monday_date, datetime.min.time(), tzinfo=CALENDAR_TIMEZONE)

    parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    if parsed.tzinfo is None:
        return parsed.replace(tzinfo=CALENDAR_TIMEZONE)
    return parsed.astimezone(CALENDAR_TIMEZONE)


def coerce_or_parse_window_et(time_min: str | None, time_max: str | None) -> tuple[datetime, datetime]:
    now_et = get_now_et()
    default_min, default_max = get_default_current_week_window_et(now_et)
    if not time_min or not time_max:
        resolved_min, resolved_max = default_min, default_max
    elif time_min.strip().lower() == "next week" or time_max.strip().lower() == "next week":
        next_monday_date = (now_et + timedelta(days=(7 - now_et.weekday()))).date()
        week_start = datetime.combine(next_monday_date, datetime.min.time(), tzinfo=CALENDAR_TIMEZONE)
        week_end_date = (week_start + timedelta(days=6)).date()
        week_end = datetime.combine(week_end_date, datetime.max.time(), tzinfo=CALENDAR_TIMEZONE).replace(microsecond=999000)
        resolved_min, resolved_max = week_start, week_end
    elif time_min.strip().lower() == "tomorrow" or time_max.strip().lower() == "tomorrow":
        tomorrow = now_et + timedelta(days=1)
        start = tomorrow.replace(hour=0, minute=0, second=0, microsecond=0)
        end = tomorrow.replace(hour=23, minute=59, second=59, microsecond=999000)
        resolved_min, resolved_max = start, end
    elif time_min.strip().lower() in {"upcoming", "this week"} or time_max.strip().lower() in {"upcoming", "this week"}:
        resolved_min, resolved_max = get_default_current_week_window_et(now_et)
    else:
        resolved_min = _parse_window_bound_to_et(time_min)
        resolved_max = _parse_window_bound_to_et(time_max)

    if resolved_max <= resolved_min:
        raise ValueError("time_max must be greater than time_min.")
    return resolved_min, resolved_max


def _extract_text_from_digital_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def _extract_text_from_docx(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text]
    return "\n".join(paragraphs).strip()


def _ocr_pdf_with_openai(path: Path, model: str = "gpt-4.1-mini", max_pages: int = 20) -> dict[str, Any]:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY is not configured.")
    client = OpenAI(api_key=api_key)

    pdf = PdfDocument(str(path))
    page_count = len(pdf)
    pages_to_process = min(page_count, max_pages)
    extracted_pages: list[dict[str, Any]] = []
    combined_parts: list[str] = []

    for index in range(pages_to_process):
        page = pdf[index]
        image = page.render(scale=2).to_pil()
        buffer = BytesIO()
        image.save(buffer, format="PNG")
        page_b64 = base64.b64encode(buffer.getvalue()).decode("ascii")
        data_url = f"data:image/png;base64,{page_b64}"
        response = client.responses.create(
            model=model,
            input=[
                {
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": "Extract all readable text from this scanned page."},
                        {"type": "input_image", "image_url": data_url},
                    ],
                }
            ],
        )
        page_text = response.output_text.strip()
        extracted_pages.append({"page_number": index + 1, "chars_extracted": len(page_text)})
        combined_parts.append(f"[Page {index + 1}]\n{page_text}")

    return {
        "text": "\n\n".join(combined_parts),
        "metadata": {
            "total_pages": page_count,
            "processed_pages": pages_to_process,
            "truncated": page_count > max_pages,
            "pages": extracted_pages,
        },
    }


def _classify_file(path: Path) -> dict[str, Any]:
    extension = path.suffix.lower()
    mime_type, _ = mimetypes.guess_type(str(path))
    mime_type = mime_type or "application/octet-stream"
    is_image = mime_type.startswith("image/")
    is_text_like = extension in {".txt", ".md", ".docx"} or mime_type.startswith("text/")
    supports_digital_pdf = extension == ".pdf"
    supports_ocr = supports_digital_pdf or is_image
    return {
        "extension": extension,
        "mime_type": mime_type,
        "is_image": is_image,
        "is_text_like": is_text_like,
        "supports_digital_pdf": supports_digital_pdf,
        "supports_ocr": supports_ocr,
    }


def _build_extraction_plan(path: Path, classification: dict[str, Any]) -> list[dict[str, Any]]:
    extension = classification["extension"]
    if extension in {".txt", ".md", ".docx"}:
        return [{"method": "native_text"}]
    if extension == ".pdf":
        return [{"method": "digital_pdf"}, {"method": "ocr_pdf"}]
    if classification["is_image"]:
        return [{"method": "ocr_image"}]
    return []


def _quality_from_text(text: str, min_chars: int = 50) -> str:
    stripped = text.strip()
    if not stripped:
        return "empty"
    if len(stripped) < min_chars:
        return "low"
    return "ok"


def _extract_via_plan(
    path: Path,
    plan: list[dict[str, Any]],
    model: str = "gpt-4.1-mini",
    max_attempts: int = 3,
) -> dict[str, Any]:
    if max_attempts < 1:
        raise ValueError("max_attempts must be at least 1.")

    attempts: list[dict[str, Any]] = []
    for step in plan[:max_attempts]:
        method = step["method"]
        try:
            if method == "native_text":
                text = extract_text_from_file(path)
                quality = _quality_from_text(text)
                attempts.append({"method": method, "success": True, "quality": quality, "chars_extracted": len(text)})
                if quality == "ok":
                    return {"text": text, "method": method, "quality": quality, "attempts": attempts}
            elif method == "digital_pdf":
                text = _extract_text_from_digital_pdf(path)
                quality = _quality_from_text(text)
                attempts.append({"method": method, "success": True, "quality": quality, "chars_extracted": len(text)})
                if quality == "ok":
                    return {"text": text, "method": method, "quality": quality, "attempts": attempts}
            elif method == "ocr_pdf":
                ocr = _ocr_pdf_with_openai(path=path, model=model)
                text = ocr["text"]
                quality = _quality_from_text(text)
                attempts.append(
                    {
                        "method": method,
                        "success": True,
                        "quality": quality,
                        "chars_extracted": len(text),
                        "ocr": ocr["metadata"],
                    }
                )
                if quality != "empty":
                    return {"text": text, "method": method, "quality": quality, "attempts": attempts}
            elif method == "ocr_image":
                raise ValueError("Image OCR extraction is not yet implemented.")
            else:
                attempts.append({"method": method, "success": False, "error": "unknown_extraction_method"})
        except Exception as exc:
            attempts.append({"method": method, "success": False, "error": str(exc)})

    return {"text": "", "method": None, "quality": "empty", "attempts": attempts, "exhausted_plan": True}


def extract_text_from_file(path: Path) -> str:
    extension = path.suffix.lower()
    if extension in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")
    if extension == ".pdf":
        return _extract_text_from_digital_pdf(path)
    if extension == ".docx":
        return _extract_text_from_docx(path)
    raise ValueError(f"Unsupported file type: {extension}")


def _summarize_text_with_openai(text: str, prompt: str | None = None, model: str = "gpt-4.1-mini") -> str:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY is not configured.")

    guidance = prompt or "Summarize the main points, key obligations, and notable risks."
    client = OpenAI(api_key=api_key)
    response = client.responses.create(
        model=model,
        input=[
            {
                "role": "system",
                "content": "You are a concise document summarization assistant.",
            },
            {
                "role": "user",
                "content": (
                    f"Guidance: {guidance}\n\n"
                    "Document text:\n"
                    f"{text[:200000]}"
                ),
            },
        ],
    )
    return response.output_text


@mcp.tool()
def weather(city: str) -> str:
    """Return current weather for a city using wttr.in."""
    response = httpx.get(f"https://wttr.in/{city}", params={"format": "j1"}, timeout=20)
    response.raise_for_status()
    data: dict[str, Any] = response.json()
    current = data["current_condition"][0]

    summary = {
        "city": city,
        "temp_c": current.get("temp_C"),
        "temp_f": current.get("temp_F"),
        "feels_like_c": current.get("FeelsLikeC"),
        "description": current.get("weatherDesc", [{}])[0].get("value"),
        "humidity": current.get("humidity"),
        "wind_kmph": current.get("windspeedKmph"),
    }
    return json.dumps(summary, indent=2)


@mcp.tool()
def list_google_calendar_events(
    calendar_id: str = "primary",
    time_min: str | None = None,
    time_max: str | None = None,
    max_results: int = 20,
) -> str:
    """List existing Google Calendar events in America/New_York within a resolved time window."""
    if max_results < 1 or max_results > 250:
        raise ValueError("max_results must be between 1 and 250.")

    service = _get_google_calendar_service()
    resolved_min_dt, resolved_max_dt = coerce_or_parse_window_et(time_min, time_max)
    resolved_min = resolved_min_dt.isoformat()
    resolved_max = resolved_max_dt.isoformat()
    response = (
        service.events()
        .list(
            calendarId=calendar_id,
            timeMin=resolved_min,
            timeMax=resolved_max,
            maxResults=max_results,
            singleEvents=True,
            orderBy="startTime",
        )
        .execute()
    )

    events = response.get("items", [])
    normalized = [
        {
            "id": event.get("id"),
            "summary": event.get("summary"),
            "description": event.get("description"),
            "location": event.get("location"),
            "status": event.get("status"),
            "start": event.get("start", {}),
            "end": event.get("end", {}),
            "html_link": event.get("htmlLink"),
        }
        for event in events
    ]
    return json.dumps(
        {
            "calendar_id": calendar_id,
            "resolved_time_min": resolved_min,
            "resolved_time_max": resolved_max,
            "resolved_timezone": "America/New_York",
            "count": len(normalized),
            "events": normalized,
        },
        indent=2,
    )


@mcp.tool()
def create_google_calendar_event(
    summary: str,
    start_iso: str,
    end_iso: str,
    calendar_id: str = "primary",
    description: str | None = None,
    location: str | None = None,
    timezone_name: str = "UTC",
) -> str:
    """Create a new Google Calendar event using RFC3339 timestamps."""
    service = _get_google_calendar_service()
    event: dict[str, Any] = {
        "summary": summary,
        "start": {"dateTime": start_iso, "timeZone": timezone_name},
        "end": {"dateTime": end_iso, "timeZone": timezone_name},
    }
    if description:
        event["description"] = description
    if location:
        event["location"] = location

    created = service.events().insert(calendarId=calendar_id, body=event).execute()
    return json.dumps(
        {
            "id": created.get("id"),
            "status": created.get("status"),
            "html_link": created.get("htmlLink"),
            "summary": created.get("summary"),
            "start": created.get("start"),
            "end": created.get("end"),
        },
        indent=2,
    )


# @mcp.tool()
# def update_google_calendar_event(
#     event_id: str,
#     summary: str | None = None,
#     description: str | None = None,
#     location: str | None = None,
#     calendar_id: str = "primary",
# ) -> str:
#     """Future option: update an existing Google Calendar event."""
#     service = _get_google_calendar_service()
#     event = service.events().get(calendarId=calendar_id, eventId=event_id).execute()
#     if summary is not None:
#         event["summary"] = summary
#     if description is not None:
#         event["description"] = description
#     if location is not None:
#         event["location"] = location
#     updated = service.events().update(calendarId=calendar_id, eventId=event_id, body=event).execute()
#     return json.dumps(updated, indent=2)
#
#
# @mcp.tool()
# def delete_google_calendar_event(event_id: str, calendar_id: str = "primary") -> str:
#     """Future option: delete an existing Google Calendar event."""
#     service = _get_google_calendar_service()
#     service.events().delete(calendarId=calendar_id, eventId=event_id).execute()
#     return json.dumps({"deleted": True, "event_id": event_id, "calendar_id": calendar_id}, indent=2)


@mcp.tool()
def query_db(sql: str) -> str:
    """Run a read-only SELECT query against local SQLite demo.db."""
    normalized = sql.strip().lower().rstrip(";")
    if not normalized.startswith("select"):
        raise ValueError("Only SELECT queries are allowed for this demo.")

    with _db_connection() as conn:
        rows = conn.execute(sql).fetchall()
    return json.dumps([dict(r) for r in rows], indent=2)


@mcp.tool()
def make_directory(path: str) -> str:
    """Create a directory inside MCP_FILE_OPS_ROOT."""
    target = _resolve_file_ops_path(path)
    target.mkdir(parents=True, exist_ok=True)
    return f"Created directory: {target}"


@mcp.tool()
def move_file(source_path: str, destination_path: str) -> str:
    """Move a file from source_path to destination_path inside MCP_FILE_OPS_ROOT."""
    source = _resolve_file_ops_path(source_path)
    destination = _resolve_file_ops_path(destination_path)

    if not source.is_file():
        raise ValueError(f"Source file does not exist: {source}")

    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.move(str(source), str(destination))
    return f"Moved file from {source} to {destination}"


@mcp.tool()
def move_files_by_glob(source_dir: str, pattern: str, destination_dir: str) -> str:
    """Move all files matching a glob pattern from source_dir into destination_dir."""
    source_root = _resolve_file_ops_path(source_dir)
    destination_root = _resolve_file_ops_path(destination_dir)

    if not source_root.is_dir():
        raise ValueError(f"Source directory does not exist: {source_root}")

    destination_root.mkdir(parents=True, exist_ok=True)

    moved: list[str] = []
    skipped: list[str] = []
    for source in sorted(source_root.glob(pattern)):
        if not source.is_file():
            continue

        destination = destination_root / source.name
        if destination.exists():
            skipped.append(source.name)
            continue

        shutil.move(str(source), str(destination))
        moved.append(source.name)

    summary = {
        "source_dir": str(source_root),
        "destination_dir": str(destination_root),
        "pattern": pattern,
        "moved_count": len(moved),
        "skipped_existing_count": len(skipped),
        "moved_files": moved,
        "skipped_existing_files": skipped,
    }
    return json.dumps(summary, indent=2)


@mcp.tool()
def list_files(path: str = ".", include_metadata: bool = False) -> str:
    """List files in a folder; set include_metadata=True for size/modified timestamps."""
    target = _resolve_file_ops_path(path)
    if not target.is_dir():
        raise ValueError(f"Not a directory: {target}")

    files = sorted(p for p in target.iterdir() if p.is_file())
    if include_metadata:
        response = [
            {
                "name": file_path.name,
                "size_bytes": file_path.stat().st_size,
                "modified_at": datetime.fromtimestamp(file_path.stat().st_mtime, tz=timezone.utc).isoformat(),
            }
            for file_path in files
        ]
    else:
        response = [p.name for p in files]
    return json.dumps(response, indent=2)


@mcp.tool()
def list_directories(path: str = ".", include_metadata: bool = False) -> str:
    """List directories in a folder; set include_metadata=True for modified timestamps."""
    target = _resolve_file_ops_path(path)
    if not target.is_dir():
        raise ValueError(f"Not a directory: {target}")

    directories = sorted(p for p in target.iterdir() if p.is_dir())
    if include_metadata:
        response = [
            {
                "name": dir_path.name,
                "modified_at": datetime.fromtimestamp(dir_path.stat().st_mtime, tz=timezone.utc).isoformat(),
            }
            for dir_path in directories
        ]
    else:
        response = [p.name for p in directories]
    return json.dumps(response, indent=2)


@mcp.tool()
def list_directory_contents(path: str = ".", include_metadata: bool = False) -> str:
    """Primary directory listing tool: return both files and directories in one response."""
    target = _resolve_file_ops_path(path)
    if not target.is_dir():
        raise ValueError(f"Not a directory: {target}")

    files = sorted(p for p in target.iterdir() if p.is_file())
    directories = sorted(p for p in target.iterdir() if p.is_dir())

    if include_metadata:
        file_entries = []
        for file_path in files:
            stat = file_path.stat()
            file_entries.append(
                {
                    "name": file_path.name,
                    "size_bytes": stat.st_size,
                    "modified_at": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
                }
            )

        directory_entries = []
        for dir_path in directories:
            stat = dir_path.stat()
            directory_entries.append(
                {
                    "name": dir_path.name,
                    "modified_at": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
                }
            )
    else:
        file_entries = [p.name for p in files]
        directory_entries = [p.name for p in directories]

    return json.dumps(
        {
            "path": str(target),
            "file_count": len(file_entries),
            "directory_count": len(directory_entries),
            "files": file_entries,
            "directories": directory_entries,
            "metadata_included": include_metadata,
            "is_empty": len(files) == 0 and len(directories) == 0,
        },
        indent=2,
    )


@mcp.tool()
def read_file(path: str) -> str:
    """Read a UTF-8 text file inside MCP_FILE_OPS_ROOT."""
    target = _resolve_file_ops_path(path)
    if not target.is_file():
        raise ValueError(f"File does not exist: {target}")
    return target.read_text(encoding="utf-8")


@mcp.tool()
def inspect_file(path: str, preview_chars: int = 4000, include_base64: bool = False) -> str:
    """Return file metadata and content preview for text/csv/image workflows."""
    target = _resolve_file_ops_path(path)
    if not target.is_file():
        raise ValueError(f"File does not exist: {target}")

    mime_type, _ = mimetypes.guess_type(str(target))
    if not mime_type:
        mime_type = "application/octet-stream"

    raw = target.read_bytes()
    result: dict[str, Any] = {
        "path": str(target),
        "name": target.name,
        "size_bytes": len(raw),
        "mime_type": mime_type,
    }

    if mime_type.startswith("text/") or mime_type in {"application/json", "text/csv"}:
        text = raw.decode("utf-8", errors="replace")
        result["text_preview"] = text[:preview_chars]
        result["text_preview_truncated"] = len(text) > preview_chars
    elif mime_type.startswith("image/"):
        result["image_note"] = "Use analyze_image_with_openai for model vision interpretation."

    if include_base64:
        result["base64"] = base64.b64encode(raw).decode("ascii")

    return json.dumps(result, indent=2)


@mcp.tool()
def analyze_image_with_openai(path: str, prompt: str, model: str = "gpt-4.1-mini") -> str:
    """Analyze an image file with an OpenAI vision-capable model."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY is not configured.")

    target = _resolve_file_ops_path(path)
    if not target.is_file():
        raise ValueError(f"File does not exist: {target}")

    mime_type, _ = mimetypes.guess_type(str(target))
    if not mime_type or not mime_type.startswith("image/"):
        raise ValueError(f"File is not an image: {target}")

    image_bytes = target.read_bytes()
    image_b64 = base64.b64encode(image_bytes).decode("ascii")
    data_url = f"data:{mime_type};base64,{image_b64}"

    client = OpenAI(api_key=api_key)
    response = client.responses.create(
        model=model,
        input=[
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": prompt},
                    {"type": "input_image", "image_url": data_url},
                ],
            }
        ],
    )
    return response.output_text


@mcp.tool()
def summarize_documents_in_folder(
    folder_path: str,
    prompt: str | None = None,
    max_files: int = 50,
    model: str = "gpt-4.1-mini",
    max_extraction_attempts: int = 3,
) -> str:
    """Summarize supported documents in a folder and return per-file + overall summaries."""
    if max_files < 1:
        raise ValueError("max_files must be at least 1.")
    folder = _resolve_file_ops_path(folder_path)
    if not folder.is_dir():
        raise ValueError(f"Not a directory: {folder}")

    files = [p for p in sorted(folder.iterdir()) if p.is_file()][:max_files]
    summaries: list[dict[str, Any]] = []
    skipped_files: list[dict[str, str]] = []

    for file_path in files:
        classification = _classify_file(file_path)
        plan = _build_extraction_plan(file_path, classification)
        if not plan:
            skipped_files.append({"path": str(file_path), "reason": "unsupported_file_type"})
            continue
        try:
            extraction = _extract_via_plan(file_path, plan, model=model, max_attempts=max_extraction_attempts)
            text = extraction["text"]
            if not text.strip():
                skipped_files.append({"path": str(file_path), "reason": "empty_or_unreadable_content"})
                continue
            summary = _summarize_text_with_openai(text=text, prompt=prompt, model=model)
            summaries.append(
                {
                    "path": str(file_path),
                    "file_name": file_path.name,
                    "summary": summary,
                    "char_count": len(text),
                    "artifact": {
                        "classification": classification,
                        "extraction_method": extraction["method"],
                        "quality": extraction["quality"],
                        "attempts": extraction["attempts"],
                    },
                }
            )
        except Exception as exc:
            skipped_files.append({"path": str(file_path), "reason": str(exc)})

    combined = "\n\n".join(
        f"{item['file_name']}:\n{item['summary']}" for item in summaries
    )[:200000]
    overall_summary = ""
    if combined:
        overall_summary = _summarize_text_with_openai(
            text=combined,
            prompt=prompt or "Create an overall summary across these file summaries.",
            model=model,
        )

    return json.dumps(
        {
            "folder_path": str(folder),
            "max_files": max_files,
            "processed_files": len(summaries),
            "skipped_files": skipped_files,
            "per_file_summaries": summaries,
            "overall_summary": overall_summary,
        },
        indent=2,
    )


@mcp.tool()
def review_contract_language(path: str, focus: str | None = None, model: str = "gpt-4.1-mini") -> str:
    """Flag potentially misleading or risky contract language (not legal advice)."""

    def _impl() -> str:
        target = _resolve_file_ops_path(path)
        if not target.is_file():
            raise ValueError(f"File does not exist: {target}")
        classification = _classify_file(target)
        plan = _build_extraction_plan(target, classification)
        if not plan:
            raise ValueError("Supported file types: .txt, .md, .pdf, .docx")

        extraction = _extract_via_plan(target, plan, model=model, max_attempts=3)
        text = extraction["text"]
        if not text.strip():
            raise ValueError("No readable text extracted from file.")

        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY is not configured.")

        guidance = focus or "Find ambiguous, one-sided, vague, or potentially misleading language."
        schema_prompt = (
            "Return strict JSON with keys: potential_issues (array), disclaimer. "
            "Each issue object must include: clause_excerpt, risk_type, severity, why_flagged, suggested_plain_language_revision."
        )
        client = OpenAI(api_key=api_key)
        response = client.responses.create(
            model=model,
            input=[
                {"role": "system", "content": "You are a contract risk reviewer. This is not legal advice."},
                {"role": "user", "content": f"{schema_prompt}\nFocus: {guidance}\n\nContract text:\n{text[:200000]}"},
            ],
        )
        output_text = response.output_text
        payload = json.loads(output_text)
        payload["document_path"] = str(target)
        payload["artifact"] = {
            "classification": classification,
            "extraction_method": extraction["method"],
            "quality": extraction["quality"],
            "attempts": extraction["attempts"],
        }
        payload["disclaimer"] = "This review is automated and is not legal advice."
        return json.dumps(payload, indent=2)

    return _run_tool_with_logging(
        tool_name="review_contract_language",
        tool_args={"path": path, "focus": focus, "model": model},
        fn=_impl,
    )


@mcp.tool()
def summarize_document(
    path: str,
    prompt: str | None = None,
    model: str = "gpt-4.1-mini",
    max_extraction_attempts: int = 3,
    output_txt_path: str | None = None,
    overwrite_output: bool = False,
) -> str:
    """Summarize one document with extraction fallback; optionally write summary to a .txt file."""
    target = _resolve_file_ops_path(path)
    if not target.is_file():
        raise ValueError(f"File does not exist: {target}")

    classification = _classify_file(target)
    plan = _build_extraction_plan(target, classification)
    if not plan:
        raise ValueError("Unsupported file type for summarization.")

    extraction = _extract_via_plan(
        target,
        plan,
        model=model,
        max_attempts=max_extraction_attempts,
    )
    text = extraction["text"]
    if not text.strip():
        raise ValueError("No readable text extracted from file after fallback attempts.")

    summary = _summarize_text_with_openai(text=text, prompt=prompt, model=model)
    result: dict[str, Any] = {
        "document_path": str(target),
        "summary": summary,
        "char_count": len(text),
        "artifact": {
            "classification": classification,
            "extraction_method": extraction["method"],
            "quality": extraction["quality"],
            "attempts": extraction["attempts"],
        },
    }

    if output_txt_path:
        output = _resolve_file_ops_path(output_txt_path)
        if output.suffix.lower() != ".txt":
            raise ValueError("output_txt_path must end with .txt")
        if output.exists() and not overwrite_output:
            raise ValueError(f"Output file exists and overwrite_output is false: {output}")
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(summary, encoding="utf-8")
        result["output_file"] = str(output)

    return json.dumps(result, indent=2)


@mcp.tool()
def extract_text_from_scanned_pdf(path: str, max_pages: int = 20, model: str = "gpt-4.1-mini") -> str:
    """Extract text from scanned/image PDFs by rendering pages and using vision."""
    if max_pages < 1:
        raise ValueError("max_pages must be at least 1.")
    target = _resolve_file_ops_path(path)
    if not target.is_file():
        raise ValueError(f"File does not exist: {target}")
    if target.suffix.lower() != ".pdf":
        raise ValueError("This OCR tool only accepts .pdf files.")

    ocr = _ocr_pdf_with_openai(path=target, model=model, max_pages=max_pages)

    return json.dumps(
        {
            "path": str(target),
            "total_pages": ocr["metadata"]["total_pages"],
            "processed_pages": ocr["metadata"]["processed_pages"],
            "truncated": ocr["metadata"]["truncated"],
            "pages": ocr["metadata"]["pages"],
            "text": ocr["text"],
        },
        indent=2,
    )


@mcp.tool()
def write_text_file(path: str, content: str, overwrite: bool = False) -> str:
    """Write a .txt file under MCP_FILE_OPS_ROOT."""
    target = _resolve_file_ops_path(path)
    if target.suffix.lower() != ".txt":
        raise ValueError("Only .txt files can be written by this tool.")
    if target.exists() and not overwrite:
        raise ValueError(f"File already exists and overwrite is false: {target}")

    target.parent.mkdir(parents=True, exist_ok=True)
    encoded = content.encode("utf-8")
    target.write_bytes(encoded)
    return json.dumps(
        {"path": str(target), "bytes_written": len(encoded), "overwrite": overwrite},
        indent=2,
    )


@mcp.tool()
def create_spreadsheet(
    path: str,
    headers: list[str] | None = None,
    overwrite: bool = False,
    delimiter: str = "comma",
) -> str:
    """Create a new empty spreadsheet (.csv/.tsv/.xlsx) under MCP_FILE_OPS_ROOT."""
    target = _resolve_file_ops_path(path)
    suffix = target.suffix.lower()
    if suffix not in SUPPORTED_SPREADSHEET_EXTENSIONS:
        raise ValueError("Only .csv, .tsv, and .xlsx files are supported.")
    if target.exists() and not overwrite:
        raise ValueError(f"File already exists and overwrite is false: {target}")

    target.parent.mkdir(parents=True, exist_ok=True)
    normalized_headers = headers or []

    if suffix == ".xlsx":
        workbook = Workbook()
        sheet = workbook.active
        if normalized_headers:
            for index, header in enumerate(normalized_headers, start=1):
                sheet.cell(row=1, column=index, value=header)
        workbook.save(target)
        return json.dumps(
            {
                "path": str(target),
                "format": "xlsx",
                "headers_written": len(normalized_headers),
                "overwrite": overwrite,
            },
            indent=2,
        )

    if delimiter not in {"comma", "tab"}:
        raise ValueError("delimiter must be either 'comma' or 'tab'.")
    separator = "," if delimiter == "comma" else "\t"
    if suffix == ".csv":
        separator = ","
    if suffix == ".tsv":
        separator = "\t"

    content = ""
    if normalized_headers:
        content = separator.join(normalized_headers) + "\n"
    target.write_text(content, encoding="utf-8")

    return json.dumps(
        {
            "path": str(target),
            "format": "csv" if separator == "," else "tsv",
            "headers_written": len(normalized_headers),
            "overwrite": overwrite,
        },
        indent=2,
    )


@mcp.tool()
def edit_spreadsheet_cell(
    path: str,
    row_index: int,
    column: str | int,
    value: str,
    has_header: bool = True,
    create_if_missing: bool = False,
) -> str:
    """Edit one cell in a .csv/.tsv/.xlsx spreadsheet under MCP_FILE_OPS_ROOT."""
    target = _resolve_file_ops_path(path)
    suffix = target.suffix.lower()
    if suffix not in SUPPORTED_SPREADSHEET_EXTENSIONS:
        raise ValueError("Only .csv, .tsv, and .xlsx files are supported.")

    if suffix == ".xlsx":
        if not target.exists():
            if not create_if_missing:
                raise ValueError(f"Spreadsheet does not exist: {target}")
            target.parent.mkdir(parents=True, exist_ok=True)
            workbook = Workbook()
            workbook.save(target)

        workbook = load_workbook(target)
        sheet = workbook.active

        if row_index < 0:
            raise ValueError("row_index must be >= 0.")

        header_row_idx = 1 if has_header else None
        data_start = 2 if has_header else 1
        target_excel_row = data_start + row_index

        if isinstance(column, int):
            if column < 0:
                raise ValueError("column (int) must be >= 0.")
            target_col_idx = column + 1
        else:
            if not has_header:
                raise ValueError("String column names require has_header=True.")

            headers: list[str] = []
            max_col = max(sheet.max_column, 1)
            for col in range(1, max_col + 1):
                cell_val = sheet.cell(row=header_row_idx, column=col).value
                headers.append("" if cell_val is None else str(cell_val))

            if column in headers:
                target_col_idx = headers.index(column) + 1
            else:
                target_col_idx = len(headers) + 1
                sheet.cell(row=header_row_idx, column=target_col_idx, value=column)

        current_value = sheet.cell(row=target_excel_row, column=target_col_idx).value
        previous_value = "" if current_value is None else str(current_value)
        sheet.cell(row=target_excel_row, column=target_col_idx, value=value)
        workbook.save(target)

        return json.dumps(
            {
                "path": str(target),
                "row_index": row_index,
                "column": column,
                "previous_value": previous_value,
                "new_value": value,
                "rows_written": sheet.max_row,
                "delimiter": "xlsx",
            },
            indent=2,
        )

    delimiter = "," if suffix == ".csv" else "\t"
    if not target.exists():
        if not create_if_missing:
            raise ValueError(f"Spreadsheet does not exist: {target}")
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("", encoding="utf-8")

    with target.open("r", encoding="utf-8", newline="") as handle:
        rows = list(csv.reader(handle, delimiter=delimiter))

    if row_index < 0:
        raise ValueError("row_index must be >= 0.")

    header: list[str] | None = rows[0] if (has_header and rows) else None
    data_start = 1 if (has_header and rows) else 0
    absolute_row = data_start + row_index

    while len(rows) <= absolute_row:
        rows.append([])

    if isinstance(column, int):
        if column < 0:
            raise ValueError("column (int) must be >= 0.")
        col_index = column
    else:
        if not has_header:
            raise ValueError("String column names require has_header=True.")
        if header is None:
            rows.insert(0, [])
            header = rows[0]
            data_start = 1
            absolute_row = data_start + row_index
            while len(rows) <= absolute_row:
                rows.append([])
        if column in header:
            col_index = header.index(column)
        else:
            header.append(column)
            col_index = len(header) - 1

    target_row = rows[absolute_row]
    while len(target_row) <= col_index:
        target_row.append("")
    previous_value = target_row[col_index]
    target_row[col_index] = value

    width = max((len(r) for r in rows), default=0)
    normalized_rows = [r + [""] * (width - len(r)) for r in rows]
    with target.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, delimiter=delimiter)
        writer.writerows(normalized_rows)

    return json.dumps(
        {
            "path": str(target),
            "row_index": row_index,
            "column": column,
            "previous_value": previous_value,
            "new_value": value,
            "rows_written": len(normalized_rows),
            "delimiter": "csv" if delimiter == "," else "tsv",
        },
        indent=2,
    )


@mcp.tool()
def append_spreadsheet_rows(
    path: str,
    rows: list[dict[str, Any] | list[str]],
    has_header: bool = True,
    create_if_missing: bool = False,
) -> str:
    """Append many rows in one operation to avoid repeated per-cell edits."""
    if not rows:
        raise ValueError("rows must include at least one row.")
    target = _resolve_file_ops_path(path)
    suffix = target.suffix.lower()
    if suffix not in SUPPORTED_SPREADSHEET_EXTENSIONS:
        raise ValueError("Only .csv, .tsv, and .xlsx files are supported.")

    if suffix == ".xlsx":
        if not target.exists():
            if not create_if_missing:
                raise ValueError(f"Spreadsheet does not exist: {target}")
            target.parent.mkdir(parents=True, exist_ok=True)
            Workbook().save(target)
        workbook = load_workbook(target)
        sheet = workbook.active
        appended = 0
        header_map: dict[str, int] = {}
        if has_header:
            max_col = max(sheet.max_column, 1)
            for col in range(1, max_col + 1):
                header_value = sheet.cell(row=1, column=col).value
                if header_value is not None and str(header_value):
                    header_map[str(header_value)] = col
        for row in rows:
            if isinstance(row, dict):
                if not has_header:
                    raise ValueError("Dict rows require has_header=True.")
                for key in row.keys():
                    if key not in header_map:
                        new_col = (max(header_map.values()) if header_map else 0) + 1
                        header_map[key] = new_col
                        sheet.cell(row=1, column=new_col, value=key)
                new_row_idx = sheet.max_row + 1
                for key, value in row.items():
                    sheet.cell(row=new_row_idx, column=header_map[key], value="" if value is None else str(value))
            else:
                list_row = ["" if v is None else str(v) for v in row]
                if has_header and header_map:
                    header_width = max(header_map.values())
                    if len(list_row) > header_width:
                        for idx in range(header_width + 1, len(list_row) + 1):
                            generated_header = f"Column {idx}"
                            header_map[generated_header] = idx
                            sheet.cell(row=1, column=idx, value=generated_header)
                        header_width = len(list_row)
                    if len(list_row) < header_width:
                        list_row.extend([""] * (header_width - len(list_row)))
                sheet.append(list_row)
            appended += 1
        workbook.save(target)
        return json.dumps({"path": str(target), "rows_appended": appended, "format": "xlsx"}, indent=2)

    delimiter = "," if suffix == ".csv" else "\t"
    if not target.exists():
        if not create_if_missing:
            raise ValueError(f"Spreadsheet does not exist: {target}")
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("", encoding="utf-8")

    with target.open("r", encoding="utf-8", newline="") as handle:
        existing_rows = list(csv.reader(handle, delimiter=delimiter))

    header: list[str] = existing_rows[0] if (has_header and existing_rows) else []
    output_rows = existing_rows[:]
    appended = 0
    for row in rows:
        if isinstance(row, dict):
            if not has_header:
                raise ValueError("Dict rows require has_header=True.")
            if not output_rows:
                output_rows.append([])
                header = output_rows[0]
            for key in row.keys():
                if key not in header:
                    header.append(key)
            values = [""] * len(header)
            for key, val in row.items():
                values[header.index(key)] = "" if val is None else str(val)
            output_rows.append(values)
        else:
            list_row = ["" if v is None else str(v) for v in row]
            if has_header and header:
                if len(list_row) > len(header):
                    for idx in range(len(header) + 1, len(list_row) + 1):
                        header.append(f"Column {idx}")
                if len(list_row) < len(header):
                    list_row.extend([""] * (len(header) - len(list_row)))
            output_rows.append(list_row)
        appended += 1

    width = max((len(r) for r in output_rows), default=0)
    normalized_rows = [r + [""] * (width - len(r)) for r in output_rows]
    with target.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, delimiter=delimiter)
        writer.writerows(normalized_rows)

    return json.dumps(
        {"path": str(target), "rows_appended": appended, "format": "csv" if delimiter == "," else "tsv"},
        indent=2,
    )


def main() -> None:
    _bootstrap_db()

    transport = MCP_TRANSPORT
    print(f"[mcp-server-demo] Bootstrapped database at: {DB_PATH.resolve()}", flush=True)

    if transport == "streamable-http":
        print(
            f"[mcp-server-demo] Starting streamable-http MCP server at "
            f"http://{MCP_HOST}:{MCP_PORT}{MCP_PATH}",
            flush=True,
        )

        if HEARTBEAT_SECONDS > 0:
            def _heartbeat() -> None:
                while True:
                    print(
                        f"[mcp-server-demo] healthy: transport=streamable-http "
                        f"url=http://{MCP_HOST}:{MCP_PORT}{MCP_PATH}",
                        flush=True,
                    )
                    time.sleep(HEARTBEAT_SECONDS)

            threading.Thread(target=_heartbeat, daemon=True).start()

        mcp.run(transport="streamable-http")
        return

    print("[mcp-server-demo] Starting stdio MCP server.", flush=True)
    print("[mcp-server-demo] Note: OpenAI HTTP client example requires streamable-http mode.", flush=True)

    mcp.run()


if __name__ == "__main__":
    main()
